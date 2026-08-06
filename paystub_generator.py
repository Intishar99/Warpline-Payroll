"""
Generates a WarpLine-branded paystub as an editable .docx file, one per agent.

Same approach as the SBG report card generator this was modeled on: build with
python-docx, return a BytesIO buffer ready to zip up and send to the browser.
"""

import os
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0E, 0x3A, 0x52)
TEAL = RGBColor(0x17, 0xA3, 0xA3)
MUTED = RGBColor(0x6B, 0x7C, 0x85)

LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "warpline_logo.png")


def _shade_cell(cell, hex_color):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_row(table, label, value, bold=False, value_color=None):
    row = table.add_row().cells
    row[0].text = label
    row[0].paragraphs[0].runs[0].font.size = Pt(10)
    row[0].paragraphs[0].runs[0].font.color.rgb = MUTED
    row[1].text = value
    run = row[1].paragraphs[0].runs[0]
    run.font.size = Pt(10)
    run.bold = bold
    if value_color:
        run.font.color.rgb = value_color
    return row


def generate_paystub_docx(agent_name, pay_period_label, payroll_result):
    """
    payroll_result: the dict returned by calc_engine.calc_agent_payroll, plus
    'name' and 'total_hours' as added by app.py's get_agent_results().
    """
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    if os.path.exists(LOGO_PATH):
        doc.add_picture(LOGO_PATH, width=Inches(1.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    title = doc.add_paragraph()
    title_run = title.add_run("Payroll Statement")
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"{agent_name}  \u2014  Pay Period: {pay_period_label}")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = MUTED

    doc.add_paragraph()

    # --- Hours & rate ---
    heading = doc.add_paragraph()
    heading.add_run("Hours & Rate").bold = True
    heading.runs[0].font.color.rgb = NAVY
    heading.runs[0].font.size = Pt(12)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(3.2)
    table.columns[1].width = Inches(2.5)

    _add_row(table, "Hours Worked", f"{payroll_result['total_hours']:.2f}")
    if payroll_result.get("manual_hours"):
        _add_row(table, "Manual Hours", f"{payroll_result['manual_hours']:.2f}")
        _add_row(table, "Effective Hours", f"{payroll_result['effective_hours']:.2f}")
    _add_row(table, "Hourly Rate", f"${payroll_result['hourly_rate']:.2f} / hr")
    _add_row(table, "Hours Pay", f"${payroll_result['hours_pay']:.2f}", bold=True)

    doc.add_paragraph()

    # --- Campaign commission breakdown ---
    if payroll_result["campaign_breakdown"]:
        heading = doc.add_paragraph()
        heading.add_run("Campaign Commission").bold = True
        heading.runs[0].font.color.rgb = NAVY
        heading.runs[0].font.size = Pt(12)

        comm_table = doc.add_table(rows=1, cols=4)
        comm_table.style = "Table Grid"
        hdr = comm_table.rows[0].cells
        for i, label in enumerate(["Division", "Valid Sits", "Sales", "Commission"]):
            hdr[i].text = label
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
            _shade_cell(hdr[i], "0E3A52")
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for c in payroll_result["campaign_breakdown"]:
            row = comm_table.add_row().cells
            row[0].text = c["division"]
            row[1].text = str(c["valid_sits"])
            row[2].text = str(c["sales"])
            row[3].text = f"${c['sit_pay'] + c['sale_pay']:.2f}"
            for cell in row:
                cell.paragraphs[0].runs[0].font.size = Pt(10)

        doc.add_paragraph()

    # --- Bonuses ---
    if payroll_result["bonuses_earned"]:
        heading = doc.add_paragraph()
        heading.add_run("Bonuses Earned").bold = True
        heading.runs[0].font.color.rgb = NAVY
        heading.runs[0].font.size = Pt(12)
        for bonus_line in payroll_result["bonuses_earned"]:
            p = doc.add_paragraph(style=None)
            run = p.add_run(f"\u2022 {bonus_line}")
            run.font.size = Pt(10)
        doc.add_paragraph()

    # --- Summary ---
    heading = doc.add_paragraph()
    heading.add_run("Summary").bold = True
    heading.runs[0].font.color.rgb = NAVY
    heading.runs[0].font.size = Pt(12)

    summary_table = doc.add_table(rows=0, cols=2)
    summary_table.style = "Table Grid"
    summary_table.columns[0].width = Inches(3.2)
    summary_table.columns[1].width = Inches(2.5)

    _add_row(summary_table, "Hours Pay", f"${payroll_result['hours_pay']:.2f}")
    _add_row(summary_table, "Commission Total", f"${payroll_result['commission_total']:.2f}")
    _add_row(summary_table, "Bonus Total", f"${payroll_result['bonus_total']:.2f}")
    _add_row(summary_table, "Spiffs", f"${payroll_result['spiffs']:.2f}")
    gross_row = _add_row(summary_table, "Gross Pay", f"${payroll_result['gross_pay']:.2f}", bold=True, value_color=TEAL)
    for cell in gross_row:
        _shade_cell(cell, "E6F6F6")
        cell.paragraphs[0].runs[0].font.size = Pt(12)

    footer = doc.add_paragraph()
    footer_run = footer.add_run("WarpLine \u2014 Your Portal to Success.")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = MUTED
    footer_run.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
